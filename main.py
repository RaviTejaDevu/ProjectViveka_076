import tkinter as tk
from tkinter import filedialog, simpledialog
import pyttsx3
import docx
import uuid

# Initializing the text-to-speech engine
engine = pyttsx3.init()

# Function to convert text to speech and save as a unique file
def text_to_speech():
    text = text_entry.get("1.0", "end-1c")  # Get the text from the entry field

    custom_filename = simpledialog.askstring("Input", "Enter a custom filename for the output audio (without extension):")

    if custom_filename:
        output_file = f"{custom_filename}.mp3"
        engine.save_to_file(text, output_file)
        engine.runAndWait()
        status_label.config(text=f"Text converted to speech and saved as {output_file}")
    else:
        status_label.config(text="Audio save canceled.")

    engine.save_to_file(text, output_file)
    engine.runAndWait()
    status_label.config(text=f"Text converted to speech and saved as {output_file}")

# Function to convert DOCX to speech and save as a unique file
def docx_to_speech():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    custom_filename = simpledialog.askstring("Input", "Enter a custom filename for the output audio (without extension):")

    if custom_filename:
        output_file = f"{custom_filename}.mp3"
        engine.save_to_file(text, output_file)
        engine.runAndWait()
        status_label.config(text=f"Text converted to speech and saved as {output_file}")
    else:
        status_label.config(text="Audio save canceled.")

    engine.save_to_file(text, output_file)
    engine.runAndWait()
    status_label.config(text=f"DOCX content converted to speech and saved as {output_file}")

# Function to save the text as a DOCX file with a custom name
def save_to_docx():
    text = text_entry.get("1.0", "end-1c")
    custom_filename = simpledialog.askstring("Input", "Enter a custom filename for the DOCX file:")
    if custom_filename:
        output_file = f"{custom_filename}.docx"
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(output_file)
        status_label.config(text=f"Text saved as {output_file}")
    else:
        status_label.config(text="File save canceled.")

# Creating the main window
root = tk.Tk()
root.title("Text-to-Speech and DOCX Converter")

# Create and configure widgets
text_label = tk.Label(root, text="Enter text:")
text_label.pack()

text_entry = tk.Text(root, height=20, width=70)
text_entry.pack()

text_to_speech_button = tk.Button(root, text="Convert Text to Speech", command=text_to_speech)
text_to_speech_button.pack()

docx_to_speech_button = tk.Button(root, text="Convert DOCX to Speech", command=docx_to_speech)
docx_to_speech_button.pack()

save_to_docx_button = tk.Button(root, text="Save Text as DOCX", command=save_to_docx)
save_to_docx_button.pack()

status_label = tk.Label(root, text="")
status_label.pack()

root.mainloop()
